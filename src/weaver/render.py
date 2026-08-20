"""Renderers: markdown template + docx styled like Leo's originals + pdf.

The docx measurements come from the seed files (Arial, 17pt name, 9.5pt section
headings, 10pt role lines with a right-aligned date tab, 9pt bullets). The pdf
is the same structure poured into a styled HTML document and printed by the
local Playwright chromium — no new dependency, no hosted service.
"""

from __future__ import annotations

import html
import re
from pathlib import Path
from typing import Any, Sequence

SECTION_SUMMARY = "Profile"
SECTION_EXPERIENCE = "Experience"
SECTION_SKILLS = "Skills"
SECTION_EDUCATION = "Education"
SECTION_AWARDS = "Awards & Recognition"
SECTION_CLIENTS = "Notable Clients"


def contact_line(structure: dict[str, Any]) -> str:
    contact = structure.get("contact") or {}
    parts = [contact.get("location"), contact.get("email"), contact.get("phone")]
    return "   |   ".join(p for p in parts if p)


def links_line(structure: dict[str, Any]) -> str:
    contact = structure.get("contact") or {}
    return "   |   ".join(contact.get("links") or [])


def _dates_and_location(entry: dict[str, Any]) -> str:
    parts = [entry.get("dates") or "", entry.get("location") or ""]
    return "  ·  ".join(p for p in parts if p)


# ---------------------------------------------------------------------- markdown


def render_markdown(structure: dict[str, Any]) -> str:
    lines: list[str] = []
    name = structure.get("name") or "Resume"
    lines.append(f"# {name}")
    if structure.get("title"):
        lines.append(f"**{structure['title']}**")
    contact = contact_line(structure)
    links = links_line(structure)
    if contact:
        lines.append("")
        lines.append(contact)
    if links:
        lines.append(links)

    if structure.get("summary"):
        lines += ["", f"## {SECTION_SUMMARY}", ""]
        for paragraph in str(structure["summary"]).split("\n\n"):
            if paragraph.strip():
                lines.append(paragraph.strip())
                lines.append("")
        lines.pop()

    experience = structure.get("experience") or []
    if experience:
        lines += ["", f"## {SECTION_EXPERIENCE}"]
        for entry in experience:
            heading = " — ".join(p for p in [entry.get("role"), entry.get("org")] if p)
            lines += ["", f"### {heading}"]
            meta = _dates_and_location(entry)
            if meta:
                lines.append(f"*{meta}*")
            if entry.get("bullets"):
                lines.append("")
                lines += [f"- {b}" for b in entry["bullets"]]

    skills = structure.get("skills") or []
    if skills:
        lines += ["", f"## {SECTION_SKILLS}", ""]
        for group in skills:
            items = ", ".join(group.get("items") or [])
            if items:
                lines.append(f"**{group.get('domain', 'Skills')}:**  {items}")
                lines.append("")
        lines.pop()

    education = structure.get("education") or []
    if education:
        lines += ["", f"## {SECTION_EDUCATION}", ""]
        for entry in education:
            if isinstance(entry, str):
                lines.append(entry)
                continue
            head = " — ".join(p for p in [entry.get("school"), entry.get("year")] if p)
            if head:
                lines.append(f"**{head}**")
            if entry.get("degree"):
                lines.append(f"{entry['degree']}")

    awards = structure.get("awards") or []
    if awards:
        lines += ["", f"## {SECTION_AWARDS}", ""]
        lines += [f"- {a}" for a in awards]

    clients = structure.get("clients") or []
    if clients:
        lines += ["", f"## {SECTION_CLIENTS}", "", ", ".join(clients) + "."]

    text = "\n".join(lines).rstrip() + "\n"
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text


def _ensure_scheme(url: str) -> str:
    value = url.strip()
    # Accept any scheme (https:, mailto:, tel: …) — not just "://" forms.
    return value if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", value) else f"https://{value}"


def _add_hyperlink(para: Any, url: str, text: str, size: float, color: Any) -> None:
    """Clickable hyperlink run (python-docx has no native add_hyperlink)."""
    import docx
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = para.part
    rel_id = part.relate_to(_ensure_scheme(url), RT.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    rFonts = docx.oxml.shared.OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    sz = docx.oxml.shared.OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    color_el = docx.oxml.shared.OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1" if color is None else str(color))
    rPr.append(color_el)
    new_run.append(rPr)
    t = docx.oxml.shared.OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


def _render_contact_paragraph(document: Any, paragraph: Any, add_run: Any,
                              contact: str, size: float, color: Any, after: float) -> None:
    """Contact line with the email rendered as a clickable mailto link."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    para = paragraph("", size=size, color=color, align=WD_ALIGN_PARAGRAPH.CENTER, after=after)
    parts = [p.strip() for p in contact.split("|")]
    for index, part in enumerate(parts):
        if index:
            add_run(para, "   |   ", size=size, color=color)
        email_match = re.match(r"^[\w.+-]+@[\w-]+\.[\w.]+$", part)
        if email_match:
            _add_hyperlink(para, f"mailto:{part}", part, size, color)
        else:
            add_run(para, part, size=size, color=color)


def _render_links_paragraph(document: Any, paragraph: Any, add_run: Any,
                            links: str, size: float, color: Any, after: float) -> None:
    """Links line with every link clickable."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    para = paragraph("", size=size, color=color, align=WD_ALIGN_PARAGRAPH.CENTER, after=after)
    items = [item.strip() for item in links.split("|") if item.strip()]
    for index, item in enumerate(items):
        if index:
            add_run(para, "   |   ", size=size, color=color)
        _add_hyperlink(para, item, item, size, color)


# -------------------------------------------------------------------------- docx


def render_docx(structure: dict[str, Any], path: Path) -> Path:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    document = docx.Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    usable = section.page_width - section.left_margin - section.right_margin

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    gray = RGBColor(0x33, 0x33, 0x33)
    mid_gray = RGBColor(0x44, 0x44, 0x44)

    def paragraph(
        text: str = "",
        size: float = 9,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
        align: Any = None,
        before: float = 0,
        after: float = 2.5,
        style: str | None = None,
    ) -> Any:
        para = document.add_paragraph(style=style) if style else document.add_paragraph()
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)
        if align is not None:
            para.alignment = align
        if text:
            add_run(para, text, size=size, bold=bold, italic=italic, color=color)
        return para

    def add_run(
        para: Any,
        text: str,
        size: float = 9,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
    ) -> Any:
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        return run

    def heading(text: str) -> None:
        paragraph(text.upper(), size=9.5, bold=True, before=10, after=5)

    # header block
    paragraph(
        structure.get("name") or "",
        size=17,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=2,
    )
    if structure.get("title"):
        paragraph(
            structure["title"],
            size=9.5,
            color=gray,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after=3,
        )
    contact = contact_line(structure)
    links = links_line(structure)
    if contact:
        _render_contact_paragraph(document, paragraph, add_run, contact, size=8.5,
                                  color=gray, after=2 if links else 10)
    if links:
        _render_links_paragraph(document, paragraph, add_run, links, size=8.5,
                                color=gray, after=10)

    if structure.get("summary"):
        heading(SECTION_SUMMARY)
        for chunk in str(structure["summary"]).split("\n\n"):
            if chunk.strip():
                paragraph(chunk.strip(), size=9, after=3.5)

    experience = structure.get("experience") or []
    if experience:
        heading(SECTION_EXPERIENCE)
        for entry in experience:
            role_para = paragraph(before=6, after=1)
            role_para.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
            add_run(role_para, entry.get("role") or "", size=10, bold=True)
            if entry.get("org"):
                add_run(role_para, "  ·  ", size=10, color=mid_gray)
                add_run(role_para, entry["org"], size=10, bold=True)
            if entry.get("dates"):
                add_run(role_para, "\t" + entry["dates"], size=9, color=mid_gray)
            if entry.get("location"):
                paragraph(entry["location"], size=8.5, italic=True, color=mid_gray, after=2)
            for bullet in entry.get("bullets") or []:
                bullet_para = _bullet_paragraph(document, paragraph)
                add_run(bullet_para, bullet, size=9)

    skills = structure.get("skills") or []
    if skills:
        heading(SECTION_SKILLS)
        for group in skills:
            items = ", ".join(group.get("items") or [])
            if not items:
                continue
            skill_para = paragraph(after=2.5)
            add_run(skill_para, f"{group.get('domain', 'Skills')}:  ", size=9, bold=True)
            add_run(skill_para, items, size=9)

    education = structure.get("education") or []
    if education:
        heading(SECTION_EDUCATION)
        for entry in education:
            if isinstance(entry, str):
                paragraph(entry, size=9)
                continue
            head_para = paragraph(after=1)
            head_para.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
            add_run(head_para, entry.get("school") or "", size=9.5, bold=True)
            if entry.get("year"):
                add_run(head_para, "\t" + str(entry["year"]), size=9, color=mid_gray)
            if entry.get("degree"):
                paragraph(entry["degree"], size=9, after=2.5)

    awards = structure.get("awards") or []
    if awards:
        heading(SECTION_AWARDS)
        for award in awards:
            award_para = _bullet_paragraph(document, paragraph)
            add_run(award_para, award, size=9)

    clients = structure.get("clients") or []
    if clients:
        heading(SECTION_CLIENTS)
        paragraph(", ".join(clients) + ".", size=9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


def _bullet_paragraph(document: Any, paragraph_factory: Any) -> Any:
    """'List Bullet' when the template has it, else a manual bullet."""
    try:
        para = document.add_paragraph(style="List Bullet")
    except KeyError:
        para = paragraph_factory(after=2.5)
        para.add_run("• ")
        return para
    from docx.shared import Pt

    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2.5)
    return para


# --------------------------------------------------------------------------- pdf

PDF_CSS = """
@page { size: A4; margin: 14mm 15mm; }
* { box-sizing: border-box; }
body {
  font-family: Georgia, "Times New Roman", "Liberation Serif", serif;
  font-size: 10.5pt;
  line-height: 1.38;
  color: #1a1a1a;
  margin: 0;
}
a { color: #0563c1; text-decoration: none; }
header { text-align: center; margin-bottom: 12pt; }
h1.name {
  font-size: 21pt;
  font-weight: 700;
  letter-spacing: 0.4pt;
  margin: 0 0 2pt 0;
}
.title { font-size: 11pt; color: #333333; margin: 0 0 3pt 0; }
.contact, .links { font-size: 9pt; color: #333333; margin: 0 0 2pt 0; }
h2.section {
  font-size: 10pt;
  font-weight: 700;
  text-transform: uppercase;
  letter-spacing: 0.8pt;
  border-bottom: 0.6pt solid #999999;
  padding-bottom: 2pt;
  margin: 13pt 0 6pt 0;
}
p { margin: 0 0 4pt 0; }
.entry { margin-bottom: 8pt; page-break-inside: avoid; }
.entry-head { display: flex; justify-content: space-between; align-items: baseline; gap: 10pt; }
.role { font-size: 11pt; font-weight: 700; }
.dates { font-size: 9.5pt; color: #444444; white-space: nowrap; }
.where { font-size: 9pt; font-style: italic; color: #444444; margin: 0 0 3pt 0; }
ul { margin: 3pt 0 0 0; padding-left: 14pt; }
li { margin-bottom: 2pt; }
.skill-domain { font-weight: 700; }
.school { font-weight: 700; }
"""


def _esc(value: Any) -> str:
    return html.escape(str(value or ""), quote=True)


def _link_html(value: str) -> str:
    item = value.strip()
    if re.match(r"^[\w.+-]+@[\w-]+\.[\w.]+$", item):
        return f'<a href="mailto:{_esc(item)}">{_esc(item)}</a>'
    return f'<a href="{_esc(_ensure_scheme(item))}">{_esc(item)}</a>'


def resume_html(structure: dict[str, Any]) -> str:
    """The same structure render_markdown reads, as a print-ready document."""
    name = structure.get("name") or "Resume"
    out: list[str] = [
        "<!DOCTYPE html>",
        '<html lang="en"><head><meta charset="utf-8">',
        f"<title>{_esc(name)}</title>",
        f"<style>{PDF_CSS}</style>",
        "</head><body>",
        "<header>",
        f'<h1 class="name">{_esc(name)}</h1>',
    ]
    if structure.get("title"):
        out.append(f'<p class="title">{_esc(structure["title"])}</p>')

    contact = structure.get("contact") or {}
    parts: list[str] = []
    if contact.get("location"):
        parts.append(_esc(contact["location"]))
    if contact.get("email"):
        parts.append(_link_html(str(contact["email"])))
    if contact.get("phone"):
        parts.append(_esc(contact["phone"]))
    if parts:
        out.append('<p class="contact">' + "&nbsp;&nbsp;|&nbsp;&nbsp;".join(parts) + "</p>")
    links = [_link_html(str(link)) for link in (contact.get("links") or []) if str(link).strip()]
    if links:
        out.append('<p class="links">' + "&nbsp;&nbsp;|&nbsp;&nbsp;".join(links) + "</p>")
    out.append("</header>")

    def section(title: str) -> None:
        out.append(f'<h2 class="section">{_esc(title)}</h2>')

    if structure.get("summary"):
        section(SECTION_SUMMARY)
        for chunk in str(structure["summary"]).split("\n\n"):
            if chunk.strip():
                out.append(f"<p>{_esc(chunk.strip())}</p>")

    experience = structure.get("experience") or []
    if experience:
        section(SECTION_EXPERIENCE)
        for entry in experience:
            out.append('<div class="entry">')
            head = _esc(entry.get("role"))
            if entry.get("org"):
                head += f"&nbsp;&nbsp;·&nbsp;&nbsp;{_esc(entry['org'])}"
            out.append('<div class="entry-head">')
            out.append(f'<span class="role">{head}</span>')
            if entry.get("dates"):
                out.append(f'<span class="dates">{_esc(entry["dates"])}</span>')
            out.append("</div>")
            if entry.get("location"):
                out.append(f'<p class="where">{_esc(entry["location"])}</p>')
            bullets = entry.get("bullets") or []
            if bullets:
                out.append("<ul>")
                out += [f"<li>{_esc(b)}</li>" for b in bullets]
                out.append("</ul>")
            out.append("</div>")

    skills = structure.get("skills") or []
    if skills:
        section(SECTION_SKILLS)
        for group in skills:
            items = ", ".join(group.get("items") or [])
            if items:
                domain = _esc(group.get("domain") or "Skills")
                out.append(f'<p><span class="skill-domain">{domain}:</span> {_esc(items)}</p>')

    education = structure.get("education") or []
    if education:
        section(SECTION_EDUCATION)
        for entry in education:
            if isinstance(entry, str):
                out.append(f"<p>{_esc(entry)}</p>")
                continue
            out.append('<div class="entry">')
            out.append('<div class="entry-head">')
            out.append(f'<span class="school">{_esc(entry.get("school"))}</span>')
            if entry.get("year"):
                out.append(f'<span class="dates">{_esc(entry["year"])}</span>')
            out.append("</div>")
            if entry.get("degree"):
                out.append(f"<p>{_esc(entry['degree'])}</p>")
            out.append("</div>")

    awards = structure.get("awards") or []
    if awards:
        section(SECTION_AWARDS)
        out.append("<ul>")
        out += [f"<li>{_esc(a)}</li>" for a in awards]
        out.append("</ul>")

    clients = structure.get("clients") or []
    if clients:
        section(SECTION_CLIENTS)
        out.append(f"<p>{_esc(', '.join(clients))}.</p>")

    out.append("</body></html>")
    return "\n".join(out)


def render_pdf(structure: dict[str, Any], path: Path) -> Path:
    """Print the styled HTML with the local Playwright chromium (no new deps)."""
    from playwright.sync_api import sync_playwright

    path.parent.mkdir(parents=True, exist_ok=True)
    document = resume_html(structure)
    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        try:
            page = browser.new_page()
            page.set_content(document, wait_until="load")
            page.pdf(path=str(path), format="A4", print_background=True)
        finally:
            browser.close()
    return path


def render(structure: dict[str, Any], path: Path, fmt: str) -> Path:
    fmt = (fmt or "md").lower()
    path.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "docx":
        return render_docx(structure, path)
    if fmt == "pdf":
        return render_pdf(structure, path)
    if fmt in ("md", "markdown", "txt"):
        path.write_text(render_markdown(structure), encoding="utf-8")
        return path
    raise ValueError(f"unsupported format: {fmt}")


def _file_slug(value: str) -> str:
    slug = "".join(ch if ch.isalnum() else "_" for ch in (value or ""))
    return re.sub(r"_{2,}", "_", slug).strip("_")


def default_filename(
    name: str,
    lens: str,
    fmt: str,
    stamp: str,
    job: dict[str, Any] | None = None,
) -> str:
    """The resume file's name — the first thing a recruiter reads.

    A job-linked tailor is named for the JOB: `Your_Name_Resume_Brand_Designer_
    Acme.docx`, not `Your_Name_design_engineer_20260816-013615.docx` — the
    name must carry the job title and company. Only the jobless lens tailor
    keeps the lens+stamp name; it has no posting to be named for.
    """
    ext = fmt if (fmt or "").lower() in ("docx", "pdf") else "md"
    slug = _file_slug(name or "resume")

    def _named(value: Any) -> str:
        # A job stored without --fetch keeps its URL as the title — slugging
        # that produced `..._Resume_https_jobs_ashbyhq_com_linear_cd5a...`
        # (morning of Aug 16). A URL is an address, never a name.
        text = str(value or "").strip()
        if "://" in text or re.match(r"(?i)(https?:|www\.)", text):
            return ""
        # ATS-generated titles sometimes arrive as "Job Application for <role>".
        # That is metadata, not useful recruiter-facing filename text.
        text = re.sub(r"(?i)^job\s+application\s+for\s+", "", text).strip()
        text = re.sub(r"(?i)^application\s+for\s+", "", text).strip()
        return _file_slug(text)

    title = _named((job or {}).get("title"))
    company = _named((job or {}).get("company"))
    if title or company:
        middle = "_".join(part for part in (title, company) if part)
        return f"{slug}_Resume_{middle}.{ext}"
    lens_slug = _file_slug(lens or "custom")
    return f"{slug}_{lens_slug}_{stamp}.{ext}"


def preview(structure: dict[str, Any], limit: int = 12) -> Sequence[str]:
    """A few lines of the markdown render, for terminal output."""
    return render_markdown(structure).splitlines()[:limit]
