"""Seed-resume ingestion: docx/pdf/md/txt -> canonical fact graph.

The deterministic parser is the default and the tested path. When an LLM key is
present `llm_structure()` gets first crack at the document and anything it
returns is validated against the same dataclasses before it reaches the DB —
so a bad LLM response degrades to the deterministic parse instead of poisoning
the graph.
"""

from __future__ import annotations

import re
import sqlite3
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable, Sequence

from . import db, llm
from .domains import CATEGORY_HINTS, classify, find_tools, normalize

SUPPORTED_SUFFIXES = {".docx", ".pdf", ".md", ".txt"}

MONTHS = (
    "jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|"
    "january|february|march|april|june|july|august|september|october|november|december"
)
_MONTH_NUM = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6, "jul": 7,
    "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}

DATE_TOKEN = rf"(?:(?:{MONTHS})\.?\s+)?(?:19|20)\d{{2}}|present|current|now|ongoing"
DATE_RANGE_RE = re.compile(
    rf"(?P<start>{DATE_TOKEN})\s*[–—\-]{{1,2}}\s*(?P<end>{DATE_TOKEN})", re.IGNORECASE
)
YEAR_RE = re.compile(r"(?:19|20)\d{2}")
SEP_RE = re.compile(r"\s{1,}[·|—–]\s{1,}|\s\|\s|\s{2,}[-]\s{2,}")
URL_RE = re.compile(
    r"^(?:https?://)?(?!\d)[a-z0-9][\w-]*(?:\.[\w-]+)*\.[a-z]{2,12}(?:/\S*)?$", re.I
)
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
PHONE_RE = re.compile(r"\+?\d[\d\s().-]{7,}\d")

METRIC_RE = re.compile(
    r"(~?\$\s?\d[\d,.]*\s*[KMB]?\b"          # $8M, ~$29K
    r"|\b\d[\d,.]*\s?%"                       # 99.9%
    r"|\b\d[\d,.]*\s?[KMB]?\+"                # 500+, 764K+, 2,000+
    r"|\b\d+x\b)",                            # 3x
    re.IGNORECASE,
)

SECTION_ALIASES = {
    "profile": "summary",
    "summary": "summary",
    "professional summary": "summary",
    "about": "summary",
    "experience": "experience",
    "work experience": "experience",
    "professional experience": "experience",
    "relevant experience": "experience",
    "earlier experience": "experience",
    "additional experience": "experience",
    "skills": "skills",
    "core skills": "skills",
    "technical skills": "skills",
    "key skills": "skills",
    "education": "education",
    "education & recognition": "education",
    "education and recognition": "education",
    "awards": "awards",
    "awards & recognition": "awards",
    "awards and recognition": "awards",
    "recognition": "awards",
    "honors": "awards",
    "notable clients": "clients",
    "clients": "clients",
    "selected clients": "clients",
    "projects": "projects",
    "selected projects": "projects",
}

# Words that start a bullet but are not an org — keeps the role-line detector honest.
_NOT_A_ROLE = re.compile(r"^(?:built|led|owned|designed|shipped|created|worked|ran|drove)\b", re.I)


@dataclass
class Line:
    text: str
    style: str = ""
    is_bullet: bool = False


@dataclass
class Role:
    title: str = ""
    org: str = ""
    start: str = ""
    end: str = ""
    location: str = ""
    links: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)


@dataclass
class ParsedDoc:
    source: str = ""
    name: str = ""
    headline: str = ""
    contacts: list[str] = field(default_factory=list)
    summaries: list[str] = field(default_factory=list)
    roles: list[Role] = field(default_factory=list)
    skill_lines: list[tuple[str, list[str]]] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    awards: list[str] = field(default_factory=list)
    clients: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


# --------------------------------------------------------------------------- read


def read_lines(path: Path) -> list[Line]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix in (".md", ".txt"):
        return _read_text(path.read_text(encoding="utf-8", errors="replace"))
    raise ValueError(f"unsupported file type: {path.suffix} ({path.name})")


def _read_docx(path: Path) -> list[Line]:
    import docx  # imported lazily so `weaver init` works without touching lxml

    document = docx.Document(str(path))
    lines: list[Line] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style is not None else ""
        is_bullet = "list" in (style or "").lower() or _has_numbering(para)
        lines.append(Line(text=text, style=style or "", is_bullet=is_bullet))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(Line(text=text, style="table", is_bullet=False))
    return lines


def _has_numbering(para: Any) -> bool:
    try:
        return para._p.pPr is not None and para._p.pPr.numPr is not None
    except AttributeError:
        return False


def _read_pdf(path: Path) -> list[Line]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return _read_text(text)


def _read_text(raw: str) -> list[Line]:
    lines: list[Line] = []
    for raw_line in raw.splitlines():
        text = raw_line.rstrip()
        if not text.strip():
            continue
        stripped = text.strip()
        is_bullet = bool(re.match(r"^[-*•‣◦]\s+", stripped))
        style = ""
        if stripped.startswith("#"):
            style = "Heading 1"
            stripped = stripped.lstrip("#").strip()
        if is_bullet:
            stripped = re.sub(r"^[-*•‣◦]\s+", "", stripped)
        stripped = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        stripped = re.sub(r"__(.+?)__", r"\1", stripped)
        lines.append(Line(text=stripped, style=style, is_bullet=is_bullet))
    return lines


# -------------------------------------------------------------------------- parse


def section_of(text: str) -> str | None:
    key = normalize(text.rstrip(":"))
    return SECTION_ALIASES.get(key)


def is_section_header(line: Line) -> bool:
    if line.is_bullet or "\t" in line.text or len(line.text) > 45:
        return False
    if section_of(line.text):
        return True
    if not line.style.lower().startswith("heading 1"):
        return False
    # A markdown '### Motion Designer · Kite (2019 – 2021)' is a role, not a section.
    if DATE_RANGE_RE.search(line.text) or SEP_RE.search(line.text):
        return False
    return len(line.text) < 45


def parse_document(lines: Sequence[Line], source: str) -> ParsedDoc:
    doc = ParsedDoc(source=source)
    section = "header"
    pending_role: Role | None = None
    header_lines: list[str] = []

    def close_role() -> None:
        nonlocal pending_role
        if pending_role is not None and (pending_role.title or pending_role.org):
            doc.roles.append(pending_role)
        pending_role = None

    for index, line in enumerate(lines):
        text = line.text.strip()
        if not text:
            continue

        if index == 0:
            # The first line is always the name, even when it is '# NAME' in markdown.
            header_lines.append(text)
            continue

        if is_section_header(line):
            close_role()
            section = section_of(text) or normalize(text)
            continue

        if section == "header":
            header_lines.append(text)
            continue

        if section == "summary":
            doc.summaries.append(_clean(text))
            continue

        if section == "experience":
            if line.is_bullet:
                if pending_role is not None:
                    pending_role.bullets.append(_clean(text))
                else:
                    doc.notes.append(_clean(text))
                continue
            role, trailing = parse_role_line(text)
            if role is not None:
                close_role()
                pending_role = role
                if trailing:
                    pending_role.bullets.append(_clean(trailing))
                continue
            if pending_role is not None:
                meta = parse_meta_line(text)
                if meta and not pending_role.location and not pending_role.bullets:
                    pending_role.location = meta.get("location", "") or pending_role.location
                    pending_role.start = pending_role.start or meta.get("start", "")
                    pending_role.end = pending_role.end or meta.get("end", "")
                    pending_role.links.extend(meta.get("links", []))
                elif len(text) >= 80:
                    pending_role.bullets.append(_clean(text))
                # else: an intra-role sub-heading ("Product and engineering") — skip
            continue

        if section == "skills":
            category, items = parse_skill_line(text)
            if items:
                doc.skill_lines.append((category, items))
            continue

        if section == "education":
            if _looks_like_award(text):
                doc.awards.extend(split_sentences(text))
            else:
                doc.education.append(_clean(text))
            continue

        if section == "awards":
            doc.awards.extend(split_sentences(_clean(text)))
            continue

        if section == "clients":
            if ":" in text and not text.strip().endswith(":"):
                doc.notes.append(_clean(text))
            else:
                doc.clients.extend(split_items(text))
            continue

        if section == "projects":
            doc.notes.append(_clean(text))
            continue

        # Unknown section: keep the text so it still reaches the guardrail corpus.
        doc.notes.append(_clean(text))

    close_role()
    _fill_header(doc, header_lines)
    return doc


def _fill_header(doc: ParsedDoc, header_lines: list[str]) -> None:
    for i, text in enumerate(header_lines):
        if i == 0 and not doc.name:
            doc.name = _titlecase_name(text)
            continue
        if EMAIL_RE.search(text) or PHONE_RE.search(text) or URL_RE.search(text.split("|")[0].strip()):
            doc.contacts.append(text)
            continue
        if not doc.headline:
            doc.headline = text
        else:
            doc.contacts.append(text)


def _titlecase_name(text: str) -> str:
    if text.isupper():
        return " ".join(part.capitalize() for part in text.split())
    return text


def parse_role_line(text: str) -> tuple[Role | None, str]:
    """Parse 'Title · Org · url\\tJan 2025 – Jan 2026' and friends.

    Returns (role, trailing_prose). Trailing prose only appears in the
    'Earlier Experience' one-liner style: 'Title | Org (2024–2025).  Led ...'.
    """
    raw = text.strip()
    if not raw or _NOT_A_ROLE.match(raw):
        return None, ""

    left, dates_text, trailing = raw, "", ""
    if "\t" in raw:
        head, _, tail = raw.partition("\t")
        left = head.strip()
        tail = tail.strip()
        if DATE_RANGE_RE.search(tail) or YEAR_RE.search(tail):
            dates_text = tail
        elif tail:
            trailing = tail
    else:
        # 'Title | Org (2024–2025).  Led ...' — find the parenthetical that is a date range.
        for paren in re.finditer(r"\(([^)]{4,44})\)", raw):
            if DATE_RANGE_RE.fullmatch(paren.group(1).strip()):
                dates_text = paren.group(1)
                left = raw[: paren.start()].strip()
                trailing = raw[paren.end() :].strip().lstrip(".·—- ").strip()
                break

    segments = [s.strip() for s in SEP_RE.split(left) if s and s.strip()]
    if len(segments) < 2:
        return None, ""
    if not dates_text and not trailing:
        # No dates anywhere: only accept it as a role when it reads like one.
        if not _looks_like_job_title(segments[0]):
            return None, ""

    # The org is the first segment that is neither a URL nor more job title
    # ('AI Lead & Project Manager — CGI Specialist | Acme Corp').
    title_parts = [segments[0]]
    org = ""
    links: list[str] = []
    last_index = len(segments) - 1
    for index, segment in enumerate(segments[1:], start=1):
        if URL_RE.match(segment):
            links.append(segment)
        elif not org and index < last_index and _looks_like_job_title(segment):
            title_parts.append(segment)
        elif not org:
            org = segment.strip(" .,")
        else:
            links.append(segment)
    title = " — ".join(p.strip(" .,") for p in title_parts)
    if not org:
        return None, ""
    if len(title) > 90 or len(org) > 70:
        return None, ""

    start, end = parse_dates(dates_text)
    role = Role(title=title, org=org, start=start, end=end, links=links)
    return role, trailing


_TITLE_WORDS = re.compile(
    r"\b(engineer|developer|designer|director|lead|founder|manager|artist|architect|"
    r"officer|head|specialist|consultant|technologist|creative|producer|principal|intern)\b",
    re.I,
)


def _looks_like_job_title(text: str) -> bool:
    return bool(_TITLE_WORDS.search(text))


def parse_meta_line(text: str) -> dict[str, Any]:
    """'January 2026 – Present | Portland, OR | example.com' or 'Remote'."""
    raw = text.strip()
    if len(raw) > 90 or raw.endswith(("!", "?")):
        return {}
    out: dict[str, Any] = {"links": []}
    segments = [s.strip() for s in re.split(r"\s*[|·]\s*", raw) if s.strip()]
    location_parts: list[str] = []
    for segment in segments:
        if DATE_RANGE_RE.search(segment):
            start, end = parse_dates(segment)
            out["start"], out["end"] = start, end
            continue
        if URL_RE.match(segment):
            out["links"].append(segment)
            continue
        location_parts.append(segment)
    if location_parts:
        out["location"] = ", ".join(location_parts)
    if not out.get("location") and not out.get("start"):
        return {}
    return out


def parse_dates(text: str) -> tuple[str, str]:
    if not text:
        return "", ""
    match = DATE_RANGE_RE.search(text)
    if match:
        return _tidy_date(match.group("start")), _tidy_date(match.group("end"))
    year = YEAR_RE.search(text)
    return (year.group(0), "") if year else ("", "")


def _tidy_date(value: str) -> str:
    value = re.sub(r"\s+", " ", (value or "").strip().strip(".,"))
    if not value:
        return ""
    if value.lower() in ("present", "current", "now", "ongoing"):
        return "Present"
    parts = value.split()
    if len(parts) == 2:
        month = parts[0].strip(".").lower()[:3]
        if month in _MONTH_NUM:
            return f"{parts[0].strip('.')[:3].capitalize()} {parts[1]}"
    return value


def sort_key_for(start: str, end: str) -> str:
    """Sortable 'YYYY-MM' from a start date; ongoing roles sort late."""
    year = YEAR_RE.search(start or "")
    if not year:
        year = YEAR_RE.search(end or "")
    y = int(year.group(0)) if year else 0
    month = 0
    token = (start or "").split()
    if token:
        m = token[0].strip(".").lower()[:3]
        month = _MONTH_NUM.get(m, 0)
    return f"{y:04d}-{month:02d}"


def parse_skill_line(text: str) -> tuple[str, list[str]]:
    if ":" in text:
        category, _, rest = text.partition(":")
        category = category.strip()
        items = split_items(rest)
        if items:
            return category, items
    items = split_items(text)
    return "", items if len(items) > 1 else []


def split_items(text: str) -> list[str]:
    """Comma-split that respects parentheses: 'fal.ai (Kling, Seedance)' stays whole."""
    out: list[str] = []
    buf: list[str] = []
    depth = 0
    for ch in text:
        if ch in "([":
            depth += 1
        elif ch in ")]":
            depth = max(0, depth - 1)
        if ch == "," and depth == 0:
            out.append("".join(buf))
            buf = []
        else:
            buf.append(ch)
    out.append("".join(buf))
    cleaned = []
    for item in out:
        value = _clean(item).strip(" .;")
        if value:
            cleaned.append(value)
    return cleaned


def split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z“\"])", text.strip())
    return [p.strip() for p in parts if p.strip()]


def _looks_like_award(text: str) -> bool:
    return bool(re.search(r"\b(award|awards|clios|adweek|winner|speaker|recognition)\b", text, re.I))


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\t", " ")).strip()


def extract_metrics(text: str) -> list[str]:
    return [_clean(m) for m in METRIC_RE.findall(text or "")]


def extract_projects(bullets: Iterable[str]) -> list[tuple[str, str]]:
    """(project_name, evidence_bullet) pairs — quoted names and 'X launch.' openers."""
    found: list[tuple[str, str]] = []
    seen: set[str] = set()
    for bullet in bullets:
        names: list[str] = []
        names.extend(re.findall(r"[“\"]([^”\"]{2,44})[”\"]", bullet))
        opener = re.match(
            r"^([A-Z][\w0-9.&' ]{2,40}?\s(?:launch|campaign|partnership|program|integration))\b",
            bullet,
        )
        if opener:
            names.append(opener.group(1))
        for name in names:
            name = name.strip().rstrip(",.;:'\"").strip()
            key = normalize(name)
            if not key or key in seen or len(key) < 3:
                continue
            seen.add(key)
            found.append((_clean(name), bullet))
    return found


# ------------------------------------------------------------------------- import


def norm_org(org: str) -> str:
    value = re.sub(r"\(.*?\)", " ", org or "")
    value = normalize(value)
    value = re.sub(r"\b(agency|studio|inc|llc|ltd|corp|co|ai|group|the)\b", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _fuzzy_fingerprint(
    conn: sqlite3.Connection, kind: str, text: str, threshold: float = 0.78
) -> str:
    """Reuse an existing fingerprint when a same-kind fact says nearly the same thing."""
    target = normalize(text)
    rows = conn.execute(
        "SELECT fingerprint, title FROM facts WHERE kind = ?", (kind,)
    ).fetchall()
    for row in rows:
        candidate = normalize(row["title"] or "")
        if not candidate:
            continue
        if candidate == target or SequenceMatcher(None, candidate, target).ratio() >= threshold:
            return row["fingerprint"]
    return f"{kind}|{target[:120]}"


def import_document(conn: sqlite3.Connection, path: Path, use_llm: bool = True) -> dict[str, Any]:
    """Parse one resume file into facts. Returns per-file counts."""
    source = path.name
    lines = read_lines(path)
    doc: ParsedDoc | None = None
    provider = "deterministic"
    if use_llm and llm.api_key():
        doc = llm_structure(lines, source)
        if doc is not None:
            provider = "openai"
    if doc is None:
        doc = parse_document(lines, source)
    counts = _persist(conn, doc)
    counts["source"] = source
    counts["provider"] = provider
    counts["profile"] = {
        "name": doc.name,
        "headline": doc.headline,
        "contacts": doc.contacts,
    }
    return counts


def _persist(conn: sqlite3.Connection, doc: ParsedDoc) -> dict[str, Any]:
    counts = {k: 0 for k in ("role", "summary", "skill", "tool", "metric", "project", "award", "client", "education")}

    for role in doc.roles:
        bullets = [b for b in (role.bullets or []) if b]
        text_blob = " ".join([role.title, role.org] + bullets)
        metrics = sorted({m for b in bullets for m in extract_metrics(b)})
        fingerprint = f"role|{norm_org(role.org)}|{sort_key_for(role.start, role.end)[:4]}"
        fact_id = db.upsert_fact(
            conn,
            {
                "kind": "role",
                "title": role.title,
                "org": role.org,
                "start": role.start,
                "end": role.end,
                "location": role.location,
                "bullets": bullets,
                "metrics": metrics,
                "tags": classify(text_blob),
                "source": doc.source,
                "aliases": [],
                "sort_key": sort_key_for(role.start, role.end),
                "fingerprint": fingerprint,
            },
        )
        counts["role"] += 1

        for metric in metrics:
            db.upsert_fact(
                conn,
                {
                    "kind": "metric",
                    "title": metric,
                    "org": role.org,
                    "bullets": [b for b in bullets if metric in b][:1],
                    "tags": classify(" ".join(b for b in bullets if metric in b)),
                    "source": doc.source,
                    "fingerprint": f"metric|{normalize(metric)}|{norm_org(role.org)}",
                },
            )
            counts["metric"] += 1

        for name, evidence in extract_projects(bullets):
            db.upsert_fact(
                conn,
                {
                    "kind": "project",
                    "title": name,
                    "org": role.org,
                    "start": role.start,
                    "end": role.end,
                    "bullets": [evidence],
                    "metrics": extract_metrics(evidence),
                    "tags": classify(f"{name} {evidence}"),
                    "source": doc.source,
                    "fingerprint": f"project|{normalize(name)}|{norm_org(role.org)}",
                },
            )
            counts["project"] += 1
        _ = fact_id

    for paragraph in doc.summaries:
        if len(paragraph) < 40:
            continue
        db.upsert_fact(
            conn,
            {
                "kind": "summary",
                "title": paragraph,
                "bullets": [paragraph],
                "metrics": extract_metrics(paragraph),
                "tags": classify(paragraph),
                "source": doc.source,
                "fingerprint": _fuzzy_fingerprint(conn, "summary", paragraph, 0.82),
            },
        )
        counts["summary"] += 1

    for category, items in doc.skill_lines:
        hint = CATEGORY_HINTS.get(normalize(category))
        category_domains = classify(f"{category} {' '.join(items)}")
        if hint:
            category_domains = [hint] + [d for d in category_domains if d != hint]
        db.upsert_fact(
            conn,
            {
                "kind": "skill",
                "title": category or "Skills",
                "bullets": items,
                "tags": category_domains,
                "source": doc.source,
                "fingerprint": f"skill|{normalize(category) or 'skills'}",
            },
        )
        counts["skill"] += 1
        for item in items:
            item_domains = classify(f"{category} {item}") or category_domains[:1]
            if hint and hint not in item_domains:
                item_domains = [hint] + item_domains
            db.upsert_skill(conn, item, item_domains[:3])
            for tool in find_tools(item):
                db.upsert_fact(
                    conn,
                    {
                        "kind": "tool",
                        "title": tool,
                        "tags": item_domains[:3],
                        "source": doc.source,
                        "fingerprint": f"tool|{normalize(tool)}",
                    },
                )
                counts["tool"] += 1

    for award in doc.awards:
        if len(award) < 8:
            continue
        db.upsert_fact(
            conn,
            {
                "kind": "award",
                "title": award,
                "metrics": extract_metrics(award),
                "tags": classify(award),
                "source": doc.source,
                "fingerprint": _fuzzy_fingerprint(conn, "award", award),
            },
        )
        counts["award"] += 1

    for client in doc.clients:
        name = client.strip(" .")
        if not name or len(name) > 60:
            continue
        db.upsert_fact(
            conn,
            {
                "kind": "client",
                "title": name,
                "tags": [],
                "source": doc.source,
                "fingerprint": f"client|{normalize(name)}",
            },
        )
        counts["client"] += 1

    for entry in _education_records(doc.education):
        db.upsert_fact(
            conn,
            {
                "kind": "education",
                "title": entry["degree"],
                "org": entry["school"],
                "end": entry["year"],
                "tags": ["design_engineering"],
                "source": doc.source,
                "fingerprint": f"education|{norm_org(entry['school']) or normalize(entry['degree'])}",
            },
        )
        counts["education"] += 1

    for note in doc.notes:
        if len(note) < 12:
            continue
        kind = "metric" if METRIC_RE.search(note) else "skill"
        db.upsert_fact(
            conn,
            {
                "kind": kind,
                "title": note[:120],
                "bullets": [note],
                "metrics": extract_metrics(note),
                "tags": classify(note),
                "source": doc.source,
                "fingerprint": f"{kind}|note|{normalize(note)[:100]}",
            },
        )
        counts[kind] += 1

    conn.commit()
    return counts


def _education_records(lines: Sequence[str]) -> list[dict[str, str]]:
    school = ""
    year = ""
    degrees: list[str] = []
    for line in lines:
        text = _clean(line)
        found_year = YEAR_RE.search(text)
        if re.search(r"\b(university|college|institute|school|academy)\b", text, re.I):
            school = re.sub(r"\s*\b(19|20)\d{2}\b\s*$", "", text).strip(" ,.—–-·|")
            if found_year:
                year = found_year.group(0)
            # 'Bachelor ... . Emily Carr University ..., 2022.' — degree lives here too
            if re.search(r"\b(bachelor|master|bdes|bsc|ba|phd|diploma)\b", text, re.I):
                head = re.split(r"(?=\b[A-Z][a-z]+\s+[A-Z][a-z]+\s+University)", text)[0]
                if head and head != text:
                    degrees.append(head.strip(" ,."))
                    school = text[len(head) :].strip(" ,.")
                    school = re.sub(r"\s*,?\s*\b(19|20)\d{2}\b\.?$", "", school).strip(" ,.")
            continue
        if re.search(r"\b(bachelor|master|bdes|bsc|ba|phd|diploma|major|minor)\b", text, re.I):
            degrees.append(text)
            if found_year and not year:
                year = found_year.group(0)
    if not school and not degrees:
        return []
    degree = max(degrees, key=len) if degrees else "Degree"
    degree = re.sub(r"\s*graduated\s+(19|20)\d{2}\.?\s*$", "", degree, flags=re.I).strip(" ,.")
    return [{"degree": degree, "school": school or "", "year": year}]


def mark_corroborated(conn: sqlite3.Connection) -> int:
    """verified = 1 when a fact is corroborated by 2+ seed documents."""
    rows = conn.execute("SELECT id, source FROM facts").fetchall()
    count = 0
    for row in rows:
        sources = [s for s in (row["source"] or "").split(",") if s]
        verified = 1 if len(sources) >= 2 else 0
        conn.execute("UPDATE facts SET verified = ? WHERE id = ?", (verified, row["id"]))
        count += verified
    conn.commit()
    return count


def import_directory(
    conn: sqlite3.Connection, directory: Path, use_llm: bool = True
) -> dict[str, Any]:
    files = sorted(
        p for p in directory.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    )
    per_file: list[dict[str, Any]] = []
    errors: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []
    skipped: list[dict[str, str]] = []
    for path in files:
        try:
            if looks_like_job_posting(path):
                skipped.append({"file": path.name, "reason": "reads as a job posting, not a resume"})
                continue
            counts = import_document(conn, path, use_llm=use_llm)
        except Exception as exc:  # a bad seed file must not abort the import
            errors.append({"file": path.name, "error": f"{type(exc).__name__}: {exc}"})
            continue
        per_file.append(counts)
        if not sum(v for v in counts.values() if isinstance(v, int)):
            warnings.append(
                {
                    "file": path.name,
                    "warning": "no extractable text — 0 facts "
                    "(scanned/vector PDF with no text layer?)",
                }
            )
    verified = mark_corroborated(conn)
    return {
        "files": [p.name for p in files],
        "imported": per_file,
        "errors": errors,
        "warnings": warnings,
        "skipped": skipped,
        "verified_facts": verified,
    }


POSTING_MARKERS = re.compile(
    r"\b(about the role|responsibilities|qualifications|what you.?ll do|requirements|"
    r"we are looking for|apply now|job description|nice to have|benefits|compensation)\b",
    re.I,
)


def looks_like_job_posting(path: Path) -> bool:
    """A postings file in the seed directory should not become career facts."""
    if path.suffix.lower() not in (".md", ".txt"):
        return False
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return False
    if len(POSTING_MARKERS.findall(text)) < 2:
        return False
    # A resume has an experience section; a posting does not.
    return not re.search(r"^\s*#*\s*(experience|work experience)\s*$", text, re.I | re.M)


def profile_from_imports(imports: Sequence[dict[str, Any]]) -> dict[str, Any]:
    """Best-effort contact block from the seed headers (most common wins)."""
    names = [i["profile"]["name"] for i in imports if i.get("profile", {}).get("name")]
    contacts = [c for i in imports for c in i.get("profile", {}).get("contacts", [])]
    headlines = [i["profile"]["headline"] for i in imports if i.get("profile", {}).get("headline")]
    blob = " | ".join(contacts)
    email = EMAIL_RE.search(blob)
    phone = PHONE_RE.search(blob)
    links: list[str] = []
    for chunk in re.split(r"\s*[|·]\s*", blob):
        chunk = chunk.strip()
        if URL_RE.match(chunk) and chunk.lower() not in [x.lower() for x in links]:
            links.append(chunk)
    location = ""
    for chunk in re.split(r"\s*[|·]\s*", blob):
        chunk = chunk.strip()
        if re.match(r"^[A-Z][A-Za-z .]+,\s*[A-Z]{2}", chunk):
            location = chunk
            break
    return {
        "name": max(set(names), key=names.count) if names else "",
        "title": headlines[0] if headlines else "",
        "email": email.group(0) if email else "",
        "phone": phone.group(0).strip() if phone else "",
        "location": location,
        "links": links,
    }


# ---------------------------------------------------------------------- llm assist

EXTRACT_SYSTEM = """You extract structured career facts from a resume.
Return ONLY JSON with this shape:
{"name": str, "headline": str, "contacts": [str], "summaries": [str],
 "roles": [{"title": str, "org": str, "start": str, "end": str, "location": str,
            "bullets": [str]}],
 "skill_lines": [{"category": str, "items": [str]}],
 "education": [str], "awards": [str], "clients": [str]}
Copy text verbatim from the resume. Never invent an organization, date, metric,
or bullet that is not present in the input."""


def llm_structure(lines: Sequence[Line], source: str) -> ParsedDoc | None:
    """LLM-assisted parse. Returns None whenever the response is unusable."""
    raw = "\n".join(("- " + ln.text) if ln.is_bullet else ln.text for ln in lines)
    data = llm.complete_json(EXTRACT_SYSTEM, raw[:60000])
    if data.get("_fallback"):
        return None
    try:
        doc = ParsedDoc(source=source)
        doc.name = str(data.get("name") or "")
        doc.headline = str(data.get("headline") or "")
        doc.contacts = [str(c) for c in data.get("contacts") or []]
        doc.summaries = [str(s) for s in data.get("summaries") or []]
        for entry in data.get("roles") or []:
            if not isinstance(entry, dict):
                continue
            role = Role(
                title=str(entry.get("title") or ""),
                org=str(entry.get("org") or ""),
                start=_tidy_date(str(entry.get("start") or "")),
                end=_tidy_date(str(entry.get("end") or "")),
                location=str(entry.get("location") or ""),
                bullets=[str(b) for b in entry.get("bullets") or []],
            )
            if role.title or role.org:
                doc.roles.append(role)
        for entry in data.get("skill_lines") or []:
            if isinstance(entry, dict) and entry.get("items"):
                doc.skill_lines.append(
                    (str(entry.get("category") or ""), [str(i) for i in entry["items"]])
                )
        doc.education = [str(e) for e in data.get("education") or []]
        doc.awards = [str(a) for a in data.get("awards") or []]
        doc.clients = [str(c) for c in data.get("clients") or []]
    except (TypeError, ValueError, AttributeError):
        return None
    return doc if doc.roles else None
